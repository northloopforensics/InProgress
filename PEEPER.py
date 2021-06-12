#python3
#script to read screenshots and search the contents for categories of data
try:
    from PIL import Image
except ImportError:
    import Image
import pytesseract
import webbrowser
import os
import fnmatch
import docx
from docx import Document
from docx.shared import Inches
import datetime
import PySimpleGUI as sg


def pep():
    now = datetime.datetime.now()
    document = Document()               #Creates docx file for report and provides table layout
    document.add_heading('PĒPR - Results Report', 0)
    p = document.add_paragraph()
    table = document.add_table(rows = 1, cols = 4)
    table.style = 'Table Grid'
    table.autofit = False
    col = table.columns[0]
    col.width=Inches(1.8)

    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = "Image"
    hdr_cells[1].text = "OCRd Text"
    hdr_cells[2].text = "Word List Hits"
    hdr_cells[3].text = 'File Name'
    

    folder = (values["SOURCE"])  #starts the ocr stuff
    screenshots = []
    for root, dirs, file in os.walk(folder):
        for filename in file:
            if filename.endswith(('.jpg', '.jfif', '.jpeg', '.png', '.JPG', '.JFIF', '.JPEG', '.PNG')):
                screenshots.append(os.path.join(root, filename))

    for f in screenshots:
        try:
            ocrd = [f, pytesseract.image_to_string(Image.open(f))]
        except TypeError:
            pass
        global file_title
        file_title = f
        escapes = ''.join([chr(char) for char in range(1, 32)])   #removes special and escape characters so the report docx doesn't choke
        translator = str.maketrans('_', ' ', escapes)
        
        if len(ocrd[1]) > 20:  #filters images with little/no text
            cleanocr = ocrd[1].replace('\n', ' ')
            global content
            content = cleanocr.translate(translator)
            ucontent = content.upper()  #converts translated content to uppercase for comparison
            print(content)
            #make lower case

            csamhit = ""
            narchit = ""
            placehit = ""
            usrhit = ""
            onetimehit = ""

            #converts txt files to upper case when reading
            if values['CSAM'] == True:                  #searches csam terms
                with open(".\\WORD LISTS\\Child Exploitation.txt") as csam:
                    csam = csam.readlines()
                    csam = map(str.upper, csam)
                    csam = list(map(lambda s: s.strip(), csam))
                    ret = any(ele in ucontent for ele in csam)
                    if ret == True:
                        csamhit = "Child Exploitation "
                    else:
                        csamhit = ""

            if values['NARC'] == True:                  #searches narcotic terms
                with open(".\\WORD LISTS\\Narcotics.txt") as narc:
                    narc = narc.readlines()
                    narc = map(str.upper, narc)
                    narc = list(map(lambda s: s.strip(), narc))
                    ret = any(ele in ucontent for ele in narc)
                    if ret == True:
                        narchit = "Narcotics "
                    else:
                        narchit = ""

            if values['PLACE'] == True:                  #searches in places.txt
                with open(".\\WORD LISTS\\Places.txt") as plc:
                    plc = plc.readlines()
                    plc = map(str.upper, plc)
                    plc = list(map(lambda s: s.strip(), plc))
                    ret = any(ele in ucontent for ele in plc)
                    if ret == True:
                        placehit = "Places "
                    else:
                        placehit = ""

            if values['USERMADE'] == True:                  #searches for user made txt list
                with open(".\\WORD LISTS\\User List.txt") as usr:
                    usr = usr.readlines()
                    usr = map(str.upper, usr)
                    usr = list(map(lambda s: s.strip(), usr))
                    ret = any(ele in ucontent for ele in usr)
                    if ret == True:
                        usrhit = "User List "
                    else:
                        usrhit = ""

            if values['USERLIST'] == True:                  #in GUI search window
                usrin = values['ONETIME'].splitlines()
                usrin = map(str.upper, usrin)
                ret= any(ele in ucontent for ele in usrin)
                if ret == True:
                    onetimehit = "One-Time Defined Terms"
                else:
                    onetimehit = ""
                                   
            dictionary = [csamhit, narchit, placehit, usrhit, onetimehit]
            para = document.add_paragraph()
            r = para.add_run()
            cells = table.add_row().cells
            im = cells[0].add_paragraph()
            r = im.add_run()
            
            try:
                r.add_picture(file_title, width=Inches(1.7))
            except:
                pass
            
            cells[1].text = content
            cells[2].text = dictionary
            cells[3].text = file_title
    
        document.save(values["OUTPUT"] + "\\PEPR REPORT.docx")

sg.theme('LightGrey3')   # Add a touch of color
# All the stuff inside your window.

layout = [  [sg.Text('PĒPR', size=(29, 1), font=('Arial Black', 20, 'bold underline'))],
            [sg.Text('Select the folder containing the screenshots and select a destination for the report:')],
            [sg.Text('Source Directory:'),sg.Input(key='SOURCE'), sg.FolderBrowse(key='SOURCE')],
            [sg.Text('Report Directory:'),sg.Input(key='OUTPUT'), sg.FolderBrowse(key='OUTPUT')],
            [sg.Text('_'*82)],
            [sg.Text('Select Word Lists for Comparison to Recovered Text')],
            [sg.Checkbox('Child Exploitation', key="CSAM"), sg.Checkbox('Narcotics', key="NARC"), sg.Checkbox('Personal Identifiers', disabled=True, key='PII'), sg.Checkbox('Places', key="PLACE"), sg.Checkbox('User List', key='USERMADE')],
            [sg.Text('_'*82)],
            [sg.Checkbox('Enable Search from One-Time Word List', key='USERLIST'), sg.Multiline(key="ONETIME", background_color="White"  ,enable_events=True, enter_submits=True, autoscroll=True, size=(35,5))],
            [sg.Text('_'*82)],
            [sg.Output(size=(79, 7), key="USELES")],
            [sg.Button('Ok'), sg.Button('Exit'), sg.Text(' '*78), sg.Text('Get Tesseract', font=('Arial', 10, 'underline'), text_color='blue', enable_events=True, key = '-LINK-'), sg.Text(' '*10), sg.Button('?', key='HELP')]]

# Create the Window
window = sg.Window('Reading screenshots, messages, and everything else...', layout, no_titlebar=False, grab_anywhere=False)
# Event Loop to process "events" and get the "values" of the inputs

while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Exit': # if user closes window or clicks cancel
        break
    if event == '-LINK-':
        webbrowser.open(r'https://github.com/UB-Mannheim/tesseract/wiki')
    elif event == 'HELP':
        sg.Popup("This tool is intended to assist investigators dealing with large volumes of screenshots or images of documents. The tool reads text from images found in the source directory and searches the text for terms or information deemed relevant by the user.",
                 "PĒPR makes use of Google's open-source Tesseract engine for optical character recognition (OCR) and requires the user install Tesseract prior to using PĒPR.",
                 "To download Tesseract you can click on the 'Get Tesseract' hyperlink on the PĒPR control panel."
                 "Word lists are located in the root of the program.",
                 "Each list is a text file that can be changed to meet the needs of the investigator. Each row is treated as a seperate search term. Please take steps to remove empty rows from the lists to avoid false positive responses.",
                 "The comparison between word lists and recovered text is not case sensitive.", "The OCR process used in the tool is not perfect and can be impacted by picture quality, image rotation, or other factors. Please review all output for accuracy."
                 "\n\n (c) 2021 North Loop Consulting, LLC")
    elif event == 'Ok':
        print('*****PEEPING IMAGES...*****')
        pep()
        window.refresh()
        print('*****PEEPING COMPLETE*****')
        os.startfile(values["OUTPUT"] + "\\PEPR REPORT.docx")
        window.refresh()
window.close()





